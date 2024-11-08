import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from keytopPyUtils.CommonUtils import CommonUtils
from keytopPyUtils.FileUtils import FileUtils

base_path = "/home/data/mdconverter/";
class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = markdown_content.replace("：\n", "：\n\n").replace("：\\*\\*\n", "：**\n\n")
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        FileUtils.create_paths(base_path)

    def to_html(self):
        return self.html_content

    def to_word(self):
        """
        将markdown文本转换成word，
        :return: 返回word文件标识
        """
        file_uuid = CommonUtils.generate_uuid();
        word_path = base_path+file_uuid+".docx"
        soup = BeautifulSoup(self.html_content, 'html.parser')
        # 创建 Word 文档
        doc = Document()

        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        soup = BeautifulSoup(self.html_content, 'html.parser')
        # 处理标题和段落
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table']):
            if tag.name == 'h1':
                # 添加一级标题
                heading = doc.add_heading(level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)
            elif tag.name == 'h2':
                # 添加二级标题
                heading = doc.add_heading(level=2)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)
            elif tag.name == 'h3':
                # 添加三级标题
                heading = doc.add_heading(level=3)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'h4':
                # 添加四级标题
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'p':
                # 添加段落
                paragraph = doc.add_paragraph(tag.get_text())
                run.font.size = Pt(12)
            elif tag.name == 'table':
                # 添加表格
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                # 创建表格
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'  # 应用内置表格样式
                table.alignment = WD_TABLE_ALIGNMENT.LEFT  # 表格居中对齐

                # 设置表头样式
                for cell in table.rows[0].cells:
                    # 确保段落和运行存在
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1  # 垂直居中
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)  # 设置背景色

                # 填充表格数据
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        # 确保段落和运行存在
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1  # 垂直居中

                # 设置每行的高度
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)

        # 保存 Word 文档
        doc.save(word_path)
        return file_uuid

    def to_pdf(self):
        file_uuid = CommonUtils.generate_uuid();
        pdf_path = base_path + file_uuid + ".pdf"
        html_text = self.html_content.replace('<table>',
                                      '<table style="width:100%; border-collapse: collapse; font-family: Arial, sans-serif;margin-left: 20px;">')
        html_text = html_text.replace('<th>',
                                      '<th style="border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #f2f2f2; color: #333;">')
        html_text = html_text.replace('<td>', '<td style="border: 1px solid #ddd; padding: 8px; text-align: left;">')
        html_text = html_text.replace('<tbody>', '<tbody style="background-color: #fff;">')
        HTML(string=html_text).write_pdf(pdf_path)